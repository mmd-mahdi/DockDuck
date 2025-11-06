import os
import logging
import re
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod
import PyPDF2
from docx import Document
import chardet

logger = logging.getLogger(__name__)


class BaseLoader(ABC):
    """Abstract base class for all document loaders"""

    @abstractmethod
    def load(self, file_path: str) -> Dict[str, Any]:
        pass

    @abstractmethod
    def supports_format(self, file_path: str) -> bool:
        pass


class PDFLoader(BaseLoader):
    """Improved loader for PDF documents with better text extraction"""

    def __init__(self):
        self.noise_patterns = [
            re.compile(r'^\s*page\s*\d+\s*$', re.IGNORECASE),
            re.compile(r'^\s*\d+\s*$'),  # Just numbers
            re.compile(r'^\.{10,}$'),  # Dotted lines
            re.compile(r'^[\-\=]+\s*$'),  # Lines of dashes/equals
        ]

    def is_noise_line(self, line: str) -> bool:
        """Identify and skip header/footer/page number lines"""
        line = line.strip()

        if len(line) < 3:  # Very short lines
            return True

        if line.isdigit():  # Just page numbers
            return True

        # Check against noise patterns
        for pattern in self.noise_patterns:
            if pattern.match(line):
                return True

        # Common header/footer indicators
        noise_indicators = [
            'copyright', 'all rights reserved', 'published by',
            'chapter', 'section', 'table of contents', 'contents',
            '......', '---------', '––––––', 'page', 'isbn'
        ]

        line_lower = line.lower()
        if any(indicator in line_lower for indicator in noise_indicators):
            if len(line) < 100:  # Only skip if it's short
                return True

        return False

    def extract_clean_text(self, page) -> str:
        """Extract clean text from PDF page with noise removal"""
        text = page.extract_text()

        if not text:
            return ""

        lines = text.split('\n')
        clean_lines = []

        for line in lines:
            clean_line = line.strip()
            if clean_line and not self.is_noise_line(clean_line):
                # Clean up excessive internal whitespace
                clean_line = re.sub(r'\s+', ' ', clean_line)
                clean_lines.append(clean_line)

        return ' '.join(clean_lines)

    def load(self, file_path: str) -> Dict[str, Any]:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                metadata = {
                    "file_type": "pdf",
                    "page_count": len(pdf_reader.pages),
                    "author": pdf_reader.metadata.get('/Author', ''),
                    "title": pdf_reader.metadata.get('/Title', ''),
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path)
                }

                meaningful_pages = 0
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = self.extract_clean_text(page)
                    if page_text and len(page_text) > 50:  # Only add meaningful pages
                        content += page_text + "\n\n"
                        meaningful_pages += 1

                metadata["meaningful_pages"] = meaningful_pages

                return {
                    "content": content.strip(),
                    "metadata": metadata
                }

        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to load PDF: {str(e)}")

    def supports_format(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')


class DOCXLoader(BaseLoader):
    """Improved loader for DOCX documents"""

    def load(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = Document(file_path)
            content = ""

            # Extract paragraphs with better filtering
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and len(text) > 5:  # Skip very short paragraphs
                    # Skip common header/footer content
                    if self.is_noise_text(text):
                        continue

                    # Check if this looks like a heading
                    if para.style.name.startswith('Heading'):
                        content += f"\n\n{text}\n"
                    else:
                        content += text + " "

            # Extract tables
            table_count = 0
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        table_content.append(" | ".join(row_cells))

                if table_content:
                    content += f"\n[Table]: " + " | ".join(table_content) + "\n"
                    table_count += 1

            metadata = {
                "file_type": "docx",
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": table_count,
                "file_path": file_path,
                "file_size": os.path.getsize(file_path)
            }

            return {
                "content": content.strip(),
                "metadata": metadata
            }

        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to load DOCX: {str(e)}")

    def is_noise_text(self, text: str) -> bool:
        """Check if text is likely header/footer noise"""
        text_lower = text.lower()
        noise_indicators = [
            'page', 'chapter', 'section', 'copyright',
            'confidential', 'draft', 'header', 'footer'
        ]
        return any(indicator in text_lower for indicator in noise_indicators)

    def supports_format(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc'))


class TXTLoader(BaseLoader):
    """Improved loader for plain text documents"""

    def load(self, file_path: str) -> Dict[str, Any]:
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'

            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                content = file.read()

            # Clean up text file specific issues
            lines = content.split('\n')
            clean_lines = []

            for line in lines:
                clean_line = line.strip()
                if clean_line and not self.is_noise_line(clean_line):
                    clean_lines.append(clean_line)

            clean_content = '\n'.join(clean_lines)

            metadata = {
                "file_type": "txt",
                "encoding": encoding,
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "line_count": len(clean_lines)
            }

            return {
                "content": clean_content,
                "metadata": metadata
            }

        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            raise ValueError(f"Failed to load text file: {str(e)}")

    def is_noise_line(self, line: str) -> bool:
        """Identify noise lines in text files"""
        if len(line) < 3:
            return True
        if line.count('-') > len(line) * 0.7:  # Mostly dashes
            return True
        if line.count('=') > len(line) * 0.7:  # Mostly equals
            return True
        if line.count('.') > len(line) * 0.5:  # Mostly dots
            return True
        return False

    def supports_format(self, file_path: str) -> bool:
        return file_path.lower().endswith('.txt')


class DocumentLoader:
    """Main document loader that handles multiple formats"""

    def __init__(self):
        self.loaders: List[BaseLoader] = [
            PDFLoader(),
            DOCXLoader(),
            TXTLoader()
        ]

    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a document from file path"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Find appropriate loader
        for loader in self.loaders:
            if loader.supports_format(file_path):
                logger.info(f"Loading {file_path} with {loader.__class__.__name__}")
                return loader.load(file_path)

        raise ValueError(f"Unsupported file format: {file_path}")

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        formats = []
        for loader in self.loaders:
            if isinstance(loader, PDFLoader):
                formats.append("PDF (.pdf)")
            elif isinstance(loader, DOCXLoader):
                formats.append("Word (.docx, .doc)")
            elif isinstance(loader, TXTLoader):
                formats.append("Text (.txt)")
        return formats